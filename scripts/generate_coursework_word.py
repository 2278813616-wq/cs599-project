from __future__ import annotations

import json
import math
import re
import textwrap
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "coursework_assets"
DIAGRAMS = ASSETS / "diagrams"
SCREENSHOTS = ASSETS / "screenshots"
TEST_LOGS = DOCS / "test_logs"
EXCALIDRAW = ASSETS / "excalidraw"
EXCALIDRAW_PREVIEWS = ASSETS / "excalidraw_previews"
REPORT_MD = DOCS / "CS599_大作业报告.md"
REPORT_DOCX = DOCS / "CS599_大作业报告.docx"

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
BODY_SIZE = Pt(12)  # 小四
H1_SIZE = Pt(22)  # 二号
H2_SIZE = Pt(15)
H3_SIZE = Pt(14)


EXCALIDRAW_SPECS = {
    "system_architecture": (
        "系统总体架构",
        ["Web UI 多轮交互", "FastAPI API", "Agent Core", "Tools Adapter", "Graph-RAG / Milvus", "外部 API 与 PDF"],
    ),
    "langgraph_state_machine": (
        "LangGraph 状态机",
        ["START", "load_memory_node", "intent_parse_node", "chef_node / gourmet_node", "END", "LangGraph 不可用时 SimpleGraph 降级"],
    ),
    "agent_state_schema": (
        "Agent State Schema",
        ["session_id / user_id", "mode / user_input", "健康与人数", "business_area_context", "selected_items / recent_eaten", "recommendations / report_path"],
    ),
    "agent_flow": (
        "Agent 交互流程",
        ["用户输入", "意图识别", "工具调用", "候选卡片", "用户确认", "PDF 导出与记忆写入"],
    ),
    "data_flow": (
        "数据流设计",
        ["前端请求", "Agent 状态", "高德 / 小红书 / Tavily", "结构化融合", "前端展示", "Milvus 足迹"],
    ),
    "graph_rag_milvus_memory_flow": (
        "Graph-RAG 与 Milvus 记忆流",
        ["Graph-RAG 健康规则", "Obsidian 知识卡", "Milvus 向量足迹", "近期吃过去重", "PDF 强信号 +1"],
    ),
    "home_cooking_pipeline": (
        "自己做链路",
        ["用户偏好与人数", "Milvus 近期足迹", "DeepSeek 规划 3 菜", "小红书详情 / Tavily 成品图", "备选菜单", "PDF 菜谱"],
    ),
    "dining_out_pipeline": (
        "出去吃链路",
        ["商圈选择", "高德餐厅 POI", "评分/距离/预算过滤", "小红书种草补充", "饭后娱乐检索", "PDF 探店"],
    ),
    "business_area_search_flow": (
        "商圈检索流程",
        ["输入商圈或地图选点", "浏览器定位起点", "高德 5km 检索", "去重合并", "评分 > 4.0", "距离 + 评分排序"],
    ),
    "xhs_integration_flow": (
        "小红书本地完整模式",
        ["aione 搜索笔记", "解析候选", "获取详情", "提取正文图片", "LLM 结构化", "缓存结果"],
    ),
    "tavily_image_fallback_flow": (
        "Tavily 图片容错",
        ["DeepSeek 已有菜谱结构", "Tavily 搜成品图", "图片 URL 过滤", "image-proxy 代理", "前端显示", "本地图兜底"],
    ),
    "gaode_xhs_restaurant_merge": (
        "高德 + 小红书餐厅融合",
        ["高德真实店名地址", "高德评分人均距离", "小红书图片", "推荐菜与种草描述", "以高德字段为准", "餐厅卡片"],
    ),
    "mcp_protocol_flow": (
        "MCP 协议调用",
        ["Codex / Client", "stdio JSON-RPC", "initialize", "tools/list", "tools/call", "query_diet_safety"],
    ),
    "pdf_export_flow": (
        "PDF 导出流程",
        ["同步已选卡片", "读取 session state", "按模式渲染内容", "ReportLab / Word 输出", "写入 Milvus 足迹"],
    ),
    "render_deployment_architecture": (
        "Render 部署架构",
        ["Browser", "Render Web Service", "FastAPI", "DeepSeek / 高德 / Tavily", "AIONE_ENABLED=0", "JSON fallback"],
    ),
    "observability_error_handling": (
        "错误处理与可观测性",
        ["API 超时", "工具失败", "图片代理失败", "本地模板兜底", "Tool Timeline", "debug JSON"],
    ),
    "memory_graph": (
        "未来商圈知识图谱",
        ["商圈", "餐厅", "小红书笔记", "饭后娱乐", "用户反馈", "导出 PDF 边权 +1"],
    ),
}


def font(path: str, size: int) -> ImageFont.FreeTypeFont:
    if Path(path).exists():
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


FONT_HEI_40 = font("C:/Windows/Fonts/simhei.ttf", 40)
FONT_HEI_28 = font("C:/Windows/Fonts/simhei.ttf", 28)
FONT_SONG_22 = font("C:/Windows/Fonts/simsun.ttc", 22)
FONT_SONG_18 = font("C:/Windows/Fonts/simsun.ttc", 18)


def wrap_text(text: str, width: int = 10) -> str:
    return "\n".join(textwrap.wrap(text, width=width, break_long_words=False, replace_whitespace=False))


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#2f3b52") -> None:
    draw.line([start, end], fill=color, width=3)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 14
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        x = end[0] + length * math.cos(angle + delta)
        y = end[1] + length * math.sin(angle + delta)
        draw.line([end, (x, y)], fill=color, width=3)


def render_preview(name: str, title: str, nodes: list[str]) -> Path:
    width, height = 1500, 900
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((20, 20, width - 20, height - 20), radius=28, outline="#111111", width=3)
    draw.text((70, 54), title, font=FONT_HEI_40, fill="#000000")
    draw.text((70, 108), "Obsidian Excalidraw 源图预览，可在源文件中继续编辑后截图替换 Word 占位", font=FONT_SONG_22, fill="#000000")

    box_w, box_h = 310, 120
    gap_x, gap_y = 80, 80
    start_x, start_y = 85, 205
    positions: list[tuple[int, int]] = []
    for i, node in enumerate(nodes):
        row = i // 3
        col = i % 3
        x = start_x + col * (box_w + gap_x)
        y = start_y + row * (box_h + gap_y)
        positions.append((x, y))
        bg = "#ffffff"
        draw.rounded_rectangle((x + 4, y + 5, x + box_w + 4, y + box_h + 5), radius=20, fill="#e5e7eb")
        if i % 4 == 1:
            draw.ellipse((x, y, x + box_w, y + box_h), fill=bg, outline="#111111", width=3)
        elif i % 4 == 2:
            points = [(x + box_w // 2, y), (x + box_w, y + box_h // 2), (x + box_w // 2, y + box_h), (x, y + box_h // 2)]
            draw.polygon(points, fill=bg, outline="#111111")
            draw.line(points + [points[0]], fill="#111111", width=3)
        else:
            draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=20 if i % 4 == 0 else 4, fill=bg, outline="#111111", width=3)
        draw.text((x + 24, y + 22), wrap_text(node, 11), font=FONT_HEI_28, fill="#000000")

    for i in range(len(positions) - 1):
        x1, y1 = positions[i]
        x2, y2 = positions[i + 1]
        if (i + 1) % 3 == 0:
            start = (x1 + box_w // 2, y1 + box_h + 10)
            end = (positions[i + 1][0] + box_w // 2, positions[i + 1][1] - 10)
        elif i % 3 == 2:
            continue
        else:
            start = (x1 + box_w + 8, y1 + box_h // 2)
            end = (x2 - 8, y2 + box_h // 2)
        draw_arrow(draw, start, end)

    output = EXCALIDRAW_PREVIEWS / f"{name}.png"
    img.save(output)
    return output


def excalidraw_element(element_id: str, element_type: str, **kwargs) -> dict:
    base = {
        "id": element_id,
        "type": element_type,
        "x": kwargs.get("x", 0),
        "y": kwargs.get("y", 0),
        "width": kwargs.get("width", 100),
        "height": kwargs.get("height", 60),
        "angle": 0,
        "strokeColor": kwargs.get("strokeColor", "#000000"),
        "backgroundColor": kwargs.get("backgroundColor", "transparent"),
        "fillStyle": "solid",
        "strokeWidth": 2,
        "strokeStyle": "solid",
        "roughness": 1,
        "opacity": 100,
        "groupIds": [],
        "frameId": None,
        "roundness": kwargs.get("roundness", {"type": 3} if element_type == "rectangle" else None),
        "seed": kwargs.get("seed", 1),
        "version": 1,
        "versionNonce": kwargs.get("seed", 1) * 17,
        "isDeleted": False,
        "boundElements": None,
        "updated": int(time.time() * 1000),
        "link": None,
        "locked": False,
    }
    if element_type == "text":
        text = kwargs.get("text", "")
        base.update({
            "text": text,
            "fontSize": kwargs.get("fontSize", 24),
            "fontFamily": 1,
            "textAlign": "left",
            "verticalAlign": "top",
            "containerId": None,
            "originalText": text,
            "lineHeight": 1.25,
            "baseline": kwargs.get("fontSize", 24),
        })
    if element_type == "diamond":
        base["roundness"] = None
    if element_type == "arrow":
        base.update({
            "points": kwargs.get("points", [[0, 0], [100, 0]]),
            "lastCommittedPoint": None,
            "startBinding": None,
            "endBinding": None,
            "startArrowhead": None,
            "endArrowhead": "arrow",
        })
    return base


def write_excalidraw(name: str, title: str, nodes: list[str]) -> Path:
    elements: list[dict] = [
        excalidraw_element(f"{name}_title", "text", x=60, y=40, width=900, height=60, text=title, fontSize=36, seed=11)
    ]
    box_w, box_h = 300, 100
    start_x, start_y = 80, 150
    positions: list[tuple[int, int]] = []
    for i, node in enumerate(nodes):
        row, col = divmod(i, 3)
        x = start_x + col * 380
        y = start_y + row * 180
        positions.append((x, y))
        shape = ["rectangle", "ellipse", "diamond", "rectangle"][i % 4]
        roundness = None if shape != "rectangle" else ({"type": 3} if i % 4 == 0 else {"type": 1})
        elements.append(excalidraw_element(
            f"{name}_box_{i}",
            shape,
            x=x,
            y=y,
            width=box_w,
            height=box_h,
            backgroundColor="#ffffff",
            strokeColor="#000000",
            roundness=roundness,
            seed=100 + i,
        ))
        elements.append(excalidraw_element(f"{name}_txt_{i}", "text", x=x + 18, y=y + 22, width=box_w - 36, height=box_h - 30, text=node, fontSize=22, seed=200 + i))
    for i in range(len(positions) - 1):
        x1, y1 = positions[i]
        x2, y2 = positions[i + 1]
        if i % 3 == 2:
            start_x_arrow = x1 + box_w / 2
            start_y_arrow = y1 + box_h + 12
            end_x_arrow = x2 + box_w / 2
            end_y_arrow = y2 - 12
        else:
            start_x_arrow = x1 + box_w
            start_y_arrow = y1 + box_h / 2
            end_x_arrow = x2 - 20
            end_y_arrow = y2 + box_h / 2
        elements.append(excalidraw_element(
            f"{name}_arrow_{i}",
            "arrow",
            x=start_x_arrow,
            y=start_y_arrow,
            width=end_x_arrow - start_x_arrow,
            height=end_y_arrow - start_y_arrow,
            points=[[0, 0], [end_x_arrow - start_x_arrow, end_y_arrow - start_y_arrow]],
            seed=300 + i,
        ))
    data = {
        "type": "excalidraw",
        "version": 2,
        "source": "https://excalidraw.com",
        "elements": elements,
        "appState": {"viewBackgroundColor": "#ffffff", "gridSize": None},
        "files": {},
    }
    output = EXCALIDRAW / f"{name}.excalidraw"
    output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    obsidian_output = EXCALIDRAW / f"{name}.excalidraw.md"
    obsidian_output.write_text(
        "---\n"
        "excalidraw-plugin: parsed\n"
        "tags: [excalidraw, cs599, superfoodie]\n"
        "---\n\n"
        "==⚠  Switch to EXCALIDRAW VIEW in Obsidian.==\n\n"
        "# Drawing\n"
        "```json\n"
        f"{json.dumps(data, ensure_ascii=False, indent=2)}\n"
        "```\n",
        encoding="utf-8",
    )
    return output


def generate_excalidraw_assets() -> dict[str, Path]:
    EXCALIDRAW.mkdir(parents=True, exist_ok=True)
    EXCALIDRAW_PREVIEWS.mkdir(parents=True, exist_ok=True)
    previews: dict[str, Path] = {}
    for name, (title, nodes) in EXCALIDRAW_SPECS.items():
        write_excalidraw(name, title, nodes)
        previews[name] = render_preview(name, title, nodes)
    (EXCALIDRAW / "README.md").write_text(
        "这些 .excalidraw 文件可直接导入 Excalidraw 或 Obsidian Excalidraw 插件继续编辑。"
        "Word 中已使用 excalidraw_previews 下的预览图；如需更手绘的效果，可在 Obsidian 中打开源文件后截图替换。\n",
        encoding="utf-8",
    )
    return previews


def set_run_font(run, name: str, size: Pt, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_style_font(style, name: str, size: Pt, bold: bool = False) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = size
    style.font.bold = bold


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)


def add_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, BODY_FONT, BODY_SIZE, bold=bold)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if not style else None
    paragraph.paragraph_format.line_spacing = 1.25


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, "Consolas", Pt(9))
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F8")
    paragraph._p.get_or_add_pPr().append(shading)


def add_picture_or_placeholder(doc: Document, image_path: Path, caption: str, width_cm: float = 14.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
    else:
        run = p.add_run(f"[图片占位：{caption}]")
        set_run_font(run, BODY_FONT, BODY_SIZE, bold=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, BODY_FONT, Pt(10))
    run.font.color.rgb = RGBColor(100, 116, 139)


def add_diagram_placeholder(doc: Document, caption: str, source: Path) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"【{caption}】\n\n请在 Obsidian Excalidraw 中打开源文件，截图后粘贴到此处。\n源文件：{source.relative_to(ROOT)}")
    set_run_font(run, BODY_FONT, BODY_SIZE, bold=True)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(10)
    doc.add_paragraph()


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text


def build_docx() -> None:
    previews = generate_excalidraw_assets()
    markdown = REPORT_MD.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)

    set_style_font(doc.styles["Normal"], BODY_FONT, BODY_SIZE)
    set_style_font(doc.styles["Heading 1"], HEADING_FONT, H1_SIZE, True)
    set_style_font(doc.styles["Heading 2"], HEADING_FONT, H2_SIZE, True)
    set_style_font(doc.styles["Heading 3"], HEADING_FONT, H3_SIZE, True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SuperFoodie 超级吃货智能助手\nCS599 大作业报告")
    set_run_font(run, HEADING_FONT, Pt(24), True)

    meta = doc.add_table(rows=8, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("课程名称", "CS599 企业级应用软件设计与开发"),
        ("项目名称", "SuperFoodie 超级吃货智能助手"),
        ("项目方向", "方向一：Agentic AI 原生开发"),
        ("学号", "（待填写学号）"),
        ("姓名", "（待填写姓名）"),
        ("专业", "计算机技术"),
        ("提交日期", "2026 年 6 月 22 日"),
        ("部署 URL", "Render Web Service，待部署后替换"),
    ]
    for row, (left, right) in zip(meta.rows, rows):
        row.cells[0].text = left
        row.cells[1].text = right
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, BODY_FONT, BODY_SIZE)

    doc.add_page_break()
    toc_title = doc.add_paragraph(style="Heading 1")
    toc_title.add_run("目录")
    for run in toc_title.runs:
        set_run_font(run, HEADING_FONT, H1_SIZE, True)
    add_toc(doc.add_paragraph())
    note = doc.add_paragraph()
    note_run = note.add_run("提示：打开 Word 后在目录处右键选择“更新域”，即可由引用目录自动生成页码。")
    set_run_font(note_run, BODY_FONT, Pt(10))
    doc.add_page_break()

    in_code = False
    code_lines: list[str] = []
    skip_table = False

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("# "):
            continue
        if line.startswith("| 字段") or line.startswith("|---") or line.startswith("| 课程名称") or line.startswith("| 项目名称") or line.startswith("| 项目方向") or line.startswith("| 学号") or line.startswith("| 姓名") or line.startswith("| 专业") or line.startswith("| 提交日期") or line.startswith("| 部署 URL"):
            skip_table = True
            continue
        if skip_table and line.startswith("|"):
            continue
        skip_table = False

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip() or line.strip() == "---":
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            caption, rel = image_match.group(1), image_match.group(2)
            path = DOCS / rel
            if "coursework_assets/diagrams/" in rel:
                key = Path(rel).stem
                preview = previews.get(key)
                source = EXCALIDRAW / f"{key}.excalidraw.md"
                if preview and preview.exists():
                    add_picture_or_placeholder(doc, preview, f"{caption}（Excalidraw 源文件：{source.relative_to(ROOT)}）")
                else:
                    add_diagram_placeholder(doc, caption, source)
            else:
                add_picture_or_placeholder(doc, path, caption)
            continue

        if line.startswith("## "):
            para = doc.add_paragraph(style="Heading 1")
            run = para.add_run(clean_inline(line[3:]))
            set_run_font(run, HEADING_FONT, H1_SIZE, True)
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(12)
            continue
        if line.startswith("### "):
            para = doc.add_paragraph(style="Heading 2")
            run = para.add_run(clean_inline(line[4:]))
            set_run_font(run, HEADING_FONT, H2_SIZE, True)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(clean_inline(line[2:]))
            set_run_font(run, BODY_FONT, BODY_SIZE)
            continue
        if line.startswith("> "):
            add_paragraph(doc, clean_inline(line[2:]))
            continue
        if line.startswith("|"):
            continue

        add_paragraph(doc, clean_inline(line))

    try:
        doc.save(REPORT_DOCX)
    except PermissionError:
        fallback = DOCS / "CS599_大作业报告_占位版.docx"
        doc.save(fallback)
        print(f"Original Word file is locked; generated fallback {fallback}")


if __name__ == "__main__":
    build_docx()
    print(f"Generated {REPORT_DOCX}")
    print(f"Generated Excalidraw sources: {EXCALIDRAW}")
    print(f"Generated Excalidraw previews: {EXCALIDRAW_PREVIEWS}")
