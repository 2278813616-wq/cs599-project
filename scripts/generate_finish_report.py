# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from textwrap import dedent
import html

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
INFO_DIR = ROOT / "finish" / "info"
IMG_DIR = ROOT / "finish" / "imgs"
GEN_DIR = IMG_DIR / "generated"
TOOLS_DIR = ROOT / "finish" / "tools"
GEN_DIR.mkdir(parents=True, exist_ok=True)


def find_source_doc() -> Path:
    candidates = [
        p
        for p in INFO_DIR.glob("*.docx")
        if not p.name.startswith("~$") and "(1)" not in p.stem
    ]
    if not candidates:
        raise FileNotFoundError("未找到原始 CS599 Word 文档")
    return max(candidates, key=lambda p: p.stat().st_mtime)


SOURCE_DOC = find_source_doc()
OUT_DOC = SOURCE_DOC.with_name(f"{SOURCE_DOC.stem}(1){SOURCE_DOC.suffix}")


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = font(34, True)
FONT_SUB = font(17)
FONT_NODE = font(18, True)
FONT_SMALL = font(14)
FONT_CODE = font(18)
FONT_CODE_SMALL = font(15)


def wrap(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        cur = ""
        for ch in paragraph:
            test = cur + ch
            width = draw.textbbox((0, 0), test, font=fnt)[2]
            if width <= max_width or not cur:
                cur = test
            else:
                lines.append(cur)
                cur = ch
        if cur:
            lines.append(cur)
    return lines


def draw_round_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    text_color: str = "#17202a",
    radius: int = 18,
    fnt=None,
):
    fnt = fnt or FONT_NODE
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    max_width = x2 - x1 - 28
    lines = wrap(draw, text, max_width, fnt)
    line_h = fnt.size + 8
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) // 2 + 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2
        draw.text((x, y), line, fill=text_color, font=fnt)
        y += line_h


def node_box(node: tuple[str, int, int, int, int]) -> tuple[int, int, int, int]:
    _, x, y, w, h = node
    return x, y, x + w, y + h


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color="#4b5563"):
    draw.line([start, end], fill=color, width=3)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 14 * direction, y2 - 8), (x2 - 14 * direction, y2 + 8)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 8, y2 - 14 * direction), (x2 + 8, y2 - 14 * direction)]
    draw.polygon(points, fill=color)


def canvas(title: str, subtitle: str, size=(1500, 900)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", size, "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, size[0], 92], fill="#111827")
    draw.text((36, 22), title, fill="#ffffff", font=FONT_TITLE)
    draw.text((38, 62), subtitle, fill="#cbd5e1", font=FONT_SUB)
    return img, draw


def save_drawio(name: str, title: str, nodes: list[tuple[str, int, int, int, int]], edges: list[tuple[int, int]]):
    cells = [
        '<mxCell id="0"/>',
        '<mxCell id="1" parent="0"/>',
    ]
    for idx, (label, x, y, w, h) in enumerate(nodes, start=2):
        cells.append(
            f'<mxCell id="{idx}" value="{html.escape(label)}" style="rounded=1;whiteSpace=wrap;html=1;fillColor=#dae8fc;strokeColor=#6c8ebf;fontStyle=1" vertex="1" parent="1">'
            f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
        )
    for eidx, (src, dst) in enumerate(edges, start=100):
        cells.append(
            f'<mxCell id="{eidx}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;endFill=1;" edge="1" parent="1" source="{src+2}" target="{dst+2}">'
            '<mxGeometry relative="1" as="geometry"/></mxCell>'
        )
    xml = (
        '<mxfile host="app.diagrams.net"><diagram name="'
        + html.escape(title)
        + '"><mxGraphModel><root>'
        + "".join(cells)
        + "</root></mxGraphModel></diagram></mxfile>"
    )
    (GEN_DIR / f"{name}.drawio").write_text(xml, encoding="utf-8")


def diagram_architecture():
    img, draw = canvas("SuperFoodie 系统总体架构", "前端体验层、FastAPI 编排层、AI/检索工具层与数据层解耦")
    nodes = [
        ("React 前端\n表单输入 / 推荐卡片 / PDF 预览", 70, 150, 330, 130),
        ("FastAPI 服务\n路由、会话状态、报告生成", 470, 150, 330, 130),
        ("Agent 编排器\n意图识别、状态机、并发任务", 870, 150, 330, 130),
        ("大模型服务\nQwen / DeepSeek 兼容接口", 90, 410, 310, 120),
        ("Graph-RAG + Milvus\n安全校验、记忆召回、相似去重", 470, 410, 330, 120),
        ("外部工具\n小红书 / Tavily / 高德 / MCP", 880, 410, 340, 120),
        ("本地报告输出\nReportLab PDF + 图片代理", 530, 660, 360, 110),
    ]
    colors = ["#dbeafe", "#dcfce7", "#fef3c7", "#fae8ff", "#e0f2fe", "#ffedd5", "#ede9fe"]
    outlines = ["#2563eb", "#16a34a", "#d97706", "#9333ea", "#0284c7", "#f97316", "#7c3aed"]
    for n, c, o in zip(nodes, colors, outlines):
        draw_round_box(draw, node_box(n), n[0], c, o)
    for s, e in [((400, 215), (470, 215)), ((800, 215), (870, 215)), ((1035, 280), (1035, 410)), ((470, 470), (400, 470)), ((870, 470), (800, 470)), ((635, 530), (710, 660))]:
        arrow(draw, s, e)
    img.save(GEN_DIR / "01_system_architecture_cn.png")
    save_drawio("01_system_architecture_cn", "系统总体架构", nodes, [(0, 1), (1, 2), (2, 4), (2, 5), (2, 3), (4, 6)])


def diagram_state_machine():
    img, draw = canvas("Agent 会话状态机", "从自然语言输入到推荐确认、详情补全和 PDF 生成")
    nodes = [
        ("开始\n用户输入需求", 95, 160, 260, 100),
        ("意图识别\n家庭做饭 / 外出就餐", 455, 160, 300, 100),
        ("安全校验\nGraph-RAG 食材禁忌", 860, 160, 300, 100),
        ("候选规划\n只生成 3 个菜名", 90, 390, 270, 100),
        ("并发补全\nXHS / LLM / Tavily", 455, 390, 300, 100),
        ("结果确认\n用户选择或调整", 860, 390, 300, 100),
        ("报告生成\n结构化 PDF", 455, 620, 300, 100),
    ]
    for i, n in enumerate(nodes):
        draw_round_box(draw, node_box(n), n[0], ["#dbeafe", "#dcfce7", "#fef3c7", "#fae8ff"][i % 4], "#334155")
    for s, e in [((355, 210), (455, 210)), ((755, 210), (860, 210)), ((1010, 260), (1010, 390)), ((860, 440), (755, 440)), ((455, 440), (360, 440)), ((605, 490), (605, 620))]:
        arrow(draw, s, e)
    draw.text((660, 545), "失败不阻断：单任务降级后继续合并", fill="#475569", font=FONT_SUB)
    img.save(GEN_DIR / "02_state_machine_cn.png")
    save_drawio("02_state_machine_cn", "Agent 会话状态机", nodes, [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (4, 6)])


def diagram_home_pipeline():
    img, draw = canvas("家庭做饭推荐链路", "小红书前置验证 + 紧凑菜名规划 + 三路并发详情补全")
    nodes = [
        ("用户偏好\n口味 / 人数 / 健康限制", 60, 160, 300, 105),
        ("并行启动\nXHS 可用性探测", 450, 135, 300, 100),
        ("并行启动\nLLM 菜名规划", 450, 295, 300, 100),
        ("3 个候选菜\nname / slot / keywords", 840, 220, 330, 110),
        ("每个候选并发补全\nXHS 详情 / LLM 兜底 / Tavily 图片", 500, 500, 520, 120),
        ("合并规则\nXHS > LLM > 本地模板\n图片：XHS/Tavily/本地", 520, 700, 480, 120),
    ]
    palette = ["#dbeafe", "#dcfce7", "#fef3c7", "#ede9fe", "#ffedd5", "#dcfce7"]
    for n, c in zip(nodes, palette):
        draw_round_box(draw, node_box(n), n[0], c, "#334155")
    for s, e in [((360, 212), (450, 185)), ((360, 212), (450, 345)), ((750, 185), (840, 250)), ((750, 345), (840, 250)), ((1005, 330), (790, 500)), ((760, 620), (760, 700))]:
        arrow(draw, s, e)
    img.save(GEN_DIR / "03_home_pipeline_cn.png")
    save_drawio("03_home_pipeline_cn", "家庭做饭推荐链路", nodes, [(0, 1), (0, 2), (1, 3), (2, 3), (3, 4), (4, 5)])


def diagram_dining_pipeline():
    img, draw = canvas("外出就餐推荐链路", "位置、预算和口味约束驱动的餐厅检索与筛选")
    nodes = [
        ("用户输入\n位置 / 场景 / 预算 / 口味", 80, 170, 320, 110),
        ("高德地图\nPOI 检索与距离计算", 510, 170, 320, 110),
        ("AI 重排\n人均、评分、交通、禁忌", 930, 170, 320, 110),
        ("结果展示\n餐厅卡片 + 地图链接", 300, 470, 340, 110),
        ("报告输出\n行程与推荐理由", 760, 470, 340, 110),
    ]
    for n, c in zip(nodes, ["#dbeafe", "#dcfce7", "#fef3c7", "#ede9fe", "#ffedd5"]):
        draw_round_box(draw, node_box(n), n[0], c, "#334155")
    for s, e in [((400, 225), (510, 225)), ((830, 225), (930, 225)), ((1090, 280), (930, 470)), ((930, 525), (640, 525))]:
        arrow(draw, s, e)
    img.save(GEN_DIR / "04_dining_pipeline_cn.png")
    save_drawio("04_dining_pipeline_cn", "外出就餐推荐链路", nodes, [(0, 1), (1, 2), (2, 3), (2, 4)])


def diagram_memory():
    img, draw = canvas("记忆检索与安全校验", "Graph-RAG 规则底座结合 Milvus 语义记忆，避免冲突并提升连续对话一致性")
    nodes = [
        ("用户画像\n忌口 / 健康限制 / 历史选择", 80, 190, 350, 120),
        ("Graph-RAG\n食材关系与安全规则", 540, 150, 350, 120),
        ("Milvus\n近邻召回与去重", 540, 350, 350, 120),
        ("校验器\n冲突检测与替代建议", 1000, 250, 330, 120),
        ("推荐输出\n安全、可解释、少重复", 560, 630, 380, 120),
    ]
    for n, c in zip(nodes, ["#dbeafe", "#dcfce7", "#e0f2fe", "#fef3c7", "#ede9fe"]):
        draw_round_box(draw, node_box(n), n[0], c, "#334155")
    for s, e in [((430, 250), (540, 210)), ((430, 250), (540, 410)), ((890, 210), (1000, 295)), ((890, 410), (1000, 320)), ((1165, 370), (750, 630))]:
        arrow(draw, s, e)
    img.save(GEN_DIR / "05_memory_rag_cn.png")
    save_drawio("05_memory_rag_cn", "记忆检索与安全校验", nodes, [(0, 1), (0, 2), (1, 3), (2, 3), (3, 4)])


def make_code_card(name: str, title: str, code: str):
    lines = dedent(code).strip("\n").splitlines()
    width = 1260
    height = 120 + max(1, len(lines)) * 29 + 40
    img = Image.new("RGB", (width, height), "#0f172a")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle([24, 24, width - 24, height - 24], radius=18, fill="#111827", outline="#334155", width=2)
    draw.text((56, 46), title, fill="#e5e7eb", font=FONT_NODE)
    draw.text((56, 78), "PlanetB 风格代码展示", fill="#93c5fd", font=FONT_SMALL)
    y = 120
    keywords = ["async", "await", "return", "if", "else", "for", "try", "except", "def", "with"]
    for i, line in enumerate(lines, start=1):
        draw.text((56, y), f"{i:>2}", fill="#64748b", font=FONT_CODE_SMALL)
        x = 96
        if line.strip().startswith("#"):
            color = "#94a3b8"
        elif any(line.lstrip().startswith(k) for k in keywords):
            color = "#a78bfa"
        elif "logger" in line or "time" in line:
            color = "#38bdf8"
        else:
            color = "#e5e7eb"
        draw.text((x, y), line, fill=color, font=FONT_CODE)
        y += 29
    img.save(GEN_DIR / name)


def generate_assets():
    diagram_architecture()
    diagram_state_machine()
    diagram_home_pipeline()
    diagram_dining_pipeline()
    diagram_memory()
    make_code_card(
        "code_parallel.png",
        "候选规划与小红书探测并行启动",
        """
        probe_task = asyncio.create_task(probe_xhs_available("家常菜 做法"))
        plan_task = asyncio.create_task(plan_home_recipe_candidates(user_profile))

        xhs_available, candidates = await asyncio.gather(
            probe_task,
            plan_task,
            return_exceptions=False,
        )
        """,
    )
    make_code_card(
        "code_xhs_probe.png",
        "小红书轻量前置验证",
        """
        async def probe_xhs_available(query: str, timeout: float = 3.0) -> bool:
            start = time.perf_counter()
            logger.info("xhs_probe.start query=%s", query)
            try:
                result = await asyncio.wait_for(aione_search(query, limit=3), timeout)
                ok = bool(result and result.items)
                logger.info("xhs_probe.done ok=%s cost=%.2fs", ok, time.perf_counter() - start)
                return ok
            except asyncio.TimeoutError:
                logger.warning("xhs_probe.timeout cost=%.2fs", time.perf_counter() - start)
                return False
        """,
    )
    make_code_card(
        "code_merge.png",
        "详情合并与降级规则",
        """
        if xhs_detail.ok:
            recipe = merge_recipe(xhs_detail, tavily_images)
            recipe.source = "xhs"
        elif llm_detail.ok:
            recipe = merge_recipe(llm_detail, tavily_images)
            recipe.source = "llm_fallback"
        else:
            recipe = build_local_template(candidate)
            recipe.image_url = pick_image(tavily_images, local_placeholder)
        """,
    )
    make_code_card(
        "code_mcp.png",
        "MCP 工具调用封装",
        """
        async with mcp_client.session("aione") as session:
            notes = await session.call_tool(
                "xhs_search",
                {"keyword": recipe.search_keywords, "limit": 5},
            )
            return normalize_xhs_recipe(notes)
        """,
    )


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_style(run, size=10.5, bold=False, color=None, font_name="Microsoft YaHei"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str, style=None, first_line=True):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_run_style(run, 10.5)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_style(run, 16 if level == 1 else 13, True, "1f2937")
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_style(r, 9, False, "64748b")


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 15.5):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "DBEAFE")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_style(run, 10, True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    set_run_style(run, 9.5)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.35
        run = p.add_run(item)
        set_run_style(run, 10.5)


def screenshot(name: str) -> Path:
    return IMG_DIR / name


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SuperFoodie 超级吃货智能助手")
    set_run_style(run, 22, True, "111827")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("CS599 企业级应用软件设计与开发期末大作业报告")
    set_run_style(run, 14, False, "475569")
    doc.add_paragraph()

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["课程方向", "方向一：Agentic AI 原生开发"],
            ["项目类型", "智能饮食推荐、外部工具调用、多模态报告生成"],
            ["关键技术", "FastAPI、React、Agent Workflow、Graph-RAG、Milvus、MCP、小红书、高德、Tavily、ReportLab"],
            ["增强内容", "章节文字扩充、中文重绘架构图、PlanetB 风格代码展示、运行截图整理"],
            ["提交日期", "2026 年 6 月 22 日"],
        ],
    )

    add_heading(doc, "一、选题背景与设计思想", 1)
    add_heading(doc, "1.1 问题定义", 2)
    add_paragraph(
        doc,
        "日常饮食决策看似简单，实际同时受到口味偏好、人数、预算、健康限制、食材安全、时间成本和信息可靠性影响。"
        "传统搜索引擎通常只返回零散网页，用户需要自己筛选菜谱、判断是否适合忌口、再寻找图片或餐厅信息；"
        "普通聊天机器人虽然可以生成文字建议，但缺少对真实内容平台、地图服务和历史偏好的稳定连接，容易产生不可执行或重复的推荐。"
    )
    add_paragraph(
        doc,
        "SuperFoodie 的设计目标是把“我想吃什么”转换为一条可落地的 Agentic 工作流：系统先理解用户意图，再调用检索、记忆、安全校验和外部工具，"
        "最后输出可确认、可追溯、可导出的饮食方案。项目覆盖家庭做饭和外出就餐两个高频场景，既能生成菜谱、步骤和图片，也能结合地点推荐餐厅。"
    )
    add_heading(doc, "1.2 设计思想", 2)
    add_bullets(
        doc,
        [
            "以 Agent 编排为核心：LLM 不直接包办全部结果，而是作为规划器、解释器和兜底生成器，与检索工具协同。",
            "以可靠性为优先级：外部平台可能超时、Cookie 可能失效、JSON 可能截断，因此每个关键节点都需要短超时、降级和结构化校验。",
            "以用户体验为导向：推荐结果不仅要“有答案”，还要有菜名、理由、步骤、图片、来源和最终 PDF，方便用户继续使用。",
            "以课程目标为映射：项目体现企业级应用中的规格文档、API 设计、服务拆分、状态管理、工具集成和可观测性。"
        ],
    )
    add_heading(doc, "1.3 技术路线", 2)
    add_paragraph(
        doc,
        "系统采用前后端分离架构。前端使用 React 承担输入表单、推荐卡片、图片预览和 PDF 下载；后端使用 FastAPI 提供会话 API 和 PDF 生成 API；"
        "Agent 层负责意图识别、状态流转、外部工具调用与结果合并；数据层由 Graph-RAG 规则库、Milvus 向量记忆和本地兜底模板共同构成。"
    )

    add_heading(doc, "二、Specs 规格文档", 1)
    add_paragraph(
        doc,
        "本项目在开发前先拆分了产品规格、架构规格和 API 规格。这样做的目的不是形式化写文档，而是把自然语言需求转化为可验证的工程边界："
        "产品规格描述用户目标和页面行为，架构规格描述模块职责和数据流，API 规格描述前后端契约和错误处理方式。"
    )
    add_table(
        doc,
        ["文档", "主要内容", "工程作用"],
        [
            ["product_spec.md", "用户场景、核心流程、输入输出、页面状态", "约束前端交互与推荐结果字段，避免实现偏离用户目标"],
            ["architecture_spec.md", "模块划分、Agent 状态机、Graph-RAG/Milvus/MCP 集成", "明确后端职责边界，降低外部工具失败对主链路的影响"],
            ["api_spec.md", "请求响应结构、会话状态、PDF 导出接口", "作为前后端联调和回归测试依据"],
            ["CS599_大作业报告.md", "项目背景、实现方案、测试与总结", "作为 Word 报告扩写的基础材料"],
        ],
    )
    add_paragraph(
        doc,
        "规格文档中最关键的约束是：推荐链路必须保持结构化输出，任何外部工具失败都不能导致整轮对话崩溃。"
        "因此实现时将 LLM 输出限制在明确 JSON schema 内，并在解析失败时提供本地模板和默认图片作为兜底。"
    )

    add_heading(doc, "三、系统架构与设计", 1)
    add_paragraph(
        doc,
        "系统整体架构分为四层：体验层、服务层、Agent 编排层和数据/工具层。体验层关注输入和展示，服务层统一处理 HTTP 会话与文件输出，"
        "Agent 编排层控制任务顺序与并发，数据/工具层负责把真实世界信息接入系统。这样的分层使得页面、模型和外部平台可以独立替换。"
    )
    add_image(doc, GEN_DIR / "01_system_architecture_cn.png", "图 1  SuperFoodie 系统总体架构")
    add_heading(doc, "3.1 Agent 状态机", 2)
    add_paragraph(
        doc,
        "Agent 状态机将一次请求拆成“输入理解、意图识别、安全校验、候选规划、详情补全、用户确认、报告生成”几个阶段。"
        "状态机的好处是每一步都有明确输入输出，可以记录耗时、定位失败点，也便于未来把部分阶段改为流式返回。"
    )
    add_image(doc, GEN_DIR / "02_state_machine_cn.png", "图 2  Agent 会话状态机")
    add_heading(doc, "3.2 记忆与安全校验", 2)
    add_paragraph(
        doc,
        "Graph-RAG 用于表达食材、禁忌和替代关系，例如过敏、忌口、宗教饮食或健康限制；Milvus 用于召回用户历史选择和相似菜谱，减少重复推荐。"
        "两者结合后，系统不仅能回答“推荐什么”，还能解释“为什么不推荐某些食材”。"
    )
    add_image(doc, GEN_DIR / "05_memory_rag_cn.png", "图 3  Graph-RAG 与 Milvus 记忆协同")

    add_heading(doc, "四、关键实现与代码展示", 1)
    add_paragraph(
        doc,
        "本次优化的核心是把原先“一次让大模型生成 3 个完整菜谱”的长链路拆成两段：第一段只生成 3 个候选菜名，第二段再并发补齐每道菜的步骤和图片。"
        "这样可以显著减少第一轮 token 数量，降低 JSON 截断风险，同时把外部搜索和 LLM 生成的等待时间重叠起来。"
    )
    add_image(doc, GEN_DIR / "03_home_pipeline_cn.png", "图 4  家庭做饭推荐链路")
    add_paragraph(
        doc,
        "请求开始时，系统并行启动小红书 Cookie 可用性探测和 LLM 菜名规划。小红书探测只做轻量搜索，短时间内能搜到候选即视为可用；"
        "这一步不抽取完整详情，避免把外部平台不稳定性放到主路径前面。"
    )
    add_image(doc, GEN_DIR / "code_parallel.png", "代码展示 1  并行启动 XHS 探测与 LLM 菜名规划", 15.8)
    add_image(doc, GEN_DIR / "code_xhs_probe.png", "代码展示 2  小红书轻量前置验证", 15.8)
    add_paragraph(
        doc,
        "拿到三个菜名后，每个候选都独立并发补全。若小红书可用，则优先用真实社区菜谱抽取做法和图片；若单个菜抽取失败，"
        "系统立即切换到 LLM 详情生成。Tavily 在这里不再承担步骤权威来源，只负责补充图片和来源链接。"
    )
    add_image(doc, GEN_DIR / "code_merge.png", "代码展示 3  XHS、LLM 与本地模板的合并规则", 15.8)
    add_paragraph(
        doc,
        "MCP 工具封装把外部能力统一为标准调用接口。对业务层来说，小红书、Tavily、高德地图都表现为可超时、可重试、可降级的工具任务，"
        "而不是散落在业务代码中的临时 HTTP 请求。"
    )
    add_image(doc, GEN_DIR / "code_mcp.png", "代码展示 4  MCP 工具调用封装", 15.8)
    add_image(doc, GEN_DIR / "04_dining_pipeline_cn.png", "图 5  外出就餐推荐链路")

    add_heading(doc, "五、测试与运行截图", 1)
    add_paragraph(
        doc,
        "测试重点覆盖三类风险：第一类是结构化输出是否稳定，避免 LLM 返回半截 JSON；第二类是外部工具失败时是否能降级；"
        "第三类是前端页面是否能完整展示推荐结果、图片和 PDF。以下截图来自项目运行过程。"
    )
    screenshots = [
        ("bfbcf5f3234eafd9e148a5df6ecf0a35.png", "图 6  首页与输入表单"),
        ("f34f3c16620eb1bdaf65ba465a638323.png", "图 7  用户输入示例"),
        ("5fc144f40f87d8c5100a8c22f27c67ba.png", "图 8  家庭做饭推荐结果"),
        ("4ca37ad51281f08b0afbd7a1740fb780.png", "图 9  菜谱详情与图片展示"),
        ("3d423dd8ae9cbd8e751c5307ffad513e.png", "图 10  外出就餐地图推荐"),
        ("8de757dc6adb74d706ff8151c7650ab1.png", "图 11  PDF 报告预览"),
    ]
    for file_name, cap in screenshots:
        add_image(doc, screenshot(file_name), cap, 14.6)
    add_table(
        doc,
        ["测试项", "结果", "说明"],
        [
            ["核心 pytest", "通过", "相关回归测试 1 passed，存在 1 个因外部条件跳过的用例"],
            ["语法检查", "通过", "chatbot.py、food_search.py 可完成 py_compile 检查"],
            ["真实接口 smoke", "通过", "POST /api/foodie/start 返回 200，并输出 3 个推荐候选"],
            ["外部工具降级", "通过", "小红书、Tavily 或模型详情单点失败时不阻断最终推荐"],
            ["PDF 导出", "通过", "推荐字段、图片代理和报告生成字段保持兼容"],
        ],
    )

    add_heading(doc, "六、部署、工具安装与扩展", 1)
    add_paragraph(
        doc,
        "后端可以通过本地 FastAPI 或 Render 类平台部署。由于在线环境通常不适合保存个人小红书 Cookie，部署版默认应把小红书视为可选增强，"
        "主路径依赖 LLM 详情生成和 Tavily 图片补充；本地开发环境如果 Cookie 有效，则可以启用 aione 小红书检索获得更真实的社区菜谱。"
    )
    add_table(
        doc,
        ["工具", "安装位置", "本报告中的用途"],
        [
            ["PlanetB Syntax Highlighter", str(TOOLS_DIR / "planetb-syntax-highlighter"), "用于第四部分代码展示的高亮风格参考，并生成深色代码卡片"],
            ["ai-drawio", str(TOOLS_DIR / "ai-drawio"), "用于重新组织 draw.io 图源，生成中文架构图与流程图"],
            ["python-docx + Pillow", "项目 venv", "用于自动化生成 Word、插入图片、绘制中文图和代码展示图"],
        ],
    )
    add_paragraph(
        doc,
        "后续扩展可以从三个方向继续：一是把后端并发阶段改造成前端卡片级流式展示，让用户先看到菜名再等待详情；"
        "二是引入更严格的菜谱结构校验，例如食材单位、步骤时长和热量估算；三是把用户确认后的选择写入长期记忆，提高下一轮推荐的个性化程度。"
    )

    add_heading(doc, "七、课程总结", 1)
    add_paragraph(
        doc,
        "本项目的主要收获在于把 Agentic AI 从“单次问答”推进到“可运行的软件系统”。在工程实现中，LLM 只是系统的一部分，"
        "真正决定可用性的因素还包括规格文档、状态机、并发控制、工具降级、日志观测、前端展示和报告导出。"
    )
    add_paragraph(
        doc,
        "通过这次迭代可以看到，企业级 AI 应用不能假设模型永远稳定，也不能假设外部平台永远可访问。"
        "更合理的做法是把模型调用拆小，把工具任务并行，把失败限制在局部，并在每个节点留下可观测日志。"
        "SuperFoodie 仍有进一步完善空间，但已经形成了从需求、规格、架构、实现、测试到文档交付的完整闭环。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附录：生成资产清单", 1)
    add_bullets(
        doc,
        [
            f"中文重绘图与 draw.io 源文件目录：{GEN_DIR}",
            f"原始运行截图目录：{IMG_DIR}",
            f"原始 Word 文件：{SOURCE_DOC}",
            f"增强版 Word 文件：{OUT_DOC}",
        ],
    )

    doc.save(OUT_DOC)


def main():
    generate_assets()
    build_report()
    print(OUT_DOC)


if __name__ == "__main__":
    main()
