# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from textwrap import dedent
import html
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
INFO_DIR = ROOT / "finish" / "info"
IMG_DIR = ROOT / "finish" / "imgs"
GEN_DIR = IMG_DIR / "generated_v2"
TOOLS_DIR = ROOT / "finish" / "tools"
GEN_DIR.mkdir(parents=True, exist_ok=True)


def source_doc() -> Path:
    docs = [
        p
        for p in INFO_DIR.glob("*.docx")
        if not p.name.startswith("~$") and re.search(r"CS599_.*报告", p.stem)
    ]
    if not docs:
        raise FileNotFoundError("未找到原始 CS599 Word 文档")
    original = [p for p in docs if "(1)" not in p.stem and "(2)" not in p.stem]
    return max(original or docs, key=lambda p: p.stat().st_mtime)


SOURCE_DOC = source_doc()
OUT_DOC = SOURCE_DOC.with_name(f"{SOURCE_DOC.stem}(2){SOURCE_DOC.suffix}")


def load_font(size: int, bold: bool = False, song: bool = False):
    candidates = []
    if song:
        candidates += [Path("C:/Windows/Fonts/simsun.ttc"), Path("C:/Windows/Fonts/simfang.ttf")]
    candidates += [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for p in candidates:
        if p.exists():
            return ImageFont.truetype(str(p), size=size)
    return ImageFont.load_default()


FT_TITLE = load_font(58, True)
FT_SUB = load_font(31)
FT_NODE = load_font(32, True)
FT_NODE_SMALL = load_font(27, True)
FT_CAP = load_font(24)
FT_CODE = load_font(30)
FT_CODE_SMALL = load_font(24)


def wrap(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt) -> list[str]:
    lines = []
    for part in text.split("\n"):
        cur = ""
        for ch in part:
            test = cur + ch
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width or not cur:
                cur = test
            else:
                lines.append(cur)
                cur = ch
        if cur:
            lines.append(cur)
    return lines or [""]


def box(draw, xy, text, fill, outline="#334155", fnt=None, text_color="#111827"):
    fnt = fnt or FT_NODE
    x, y, w, h = xy
    rect = [x, y, x + w, y + h]
    draw.rounded_rectangle(rect, radius=28, fill=fill, outline=outline, width=5)
    lines = wrap(draw, text, w - 56, fnt)
    line_h = fnt.size + 12
    start_y = y + (h - line_h * len(lines)) // 2
    for i, line in enumerate(lines):
        tw = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x + (w - tw) // 2, start_y + i * line_h), line, fill=text_color, font=fnt)


def arrow(draw, start, end, color="#475569"):
    draw.line([start, end], fill=color, width=6)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        d = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - 28 * d, y2 - 16), (x2 - 28 * d, y2 + 16)]
    else:
        d = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 16, y2 - 28 * d), (x2 + 16, y2 - 28 * d)]
    draw.polygon(pts, fill=color)


def base_canvas(title: str, sub: str, size=(2600, 1500)):
    img = Image.new("RGB", size, "#f8fafc")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, size[0], 165], fill="#111827")
    d.text((64, 38), title, fill="#ffffff", font=FT_TITLE)
    d.text((66, 104), sub, fill="#cbd5e1", font=FT_SUB)
    return img, d


def save_drawio(name: str, title: str, nodes: list[tuple[str, int, int, int, int]], edges: list[tuple[int, int]]):
    cells = ['<mxCell id="0"/>', '<mxCell id="1" parent="0"/>']
    for idx, (label, x, y, w, h) in enumerate(nodes, start=2):
        cells.append(
            f'<mxCell id="{idx}" value="{html.escape(label)}" style="rounded=1;whiteSpace=wrap;html=1;fillColor=#dae8fc;strokeColor=#6c8ebf;fontStyle=1" vertex="1" parent="1">'
            f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
        )
    for eidx, (src, dst) in enumerate(edges, start=100):
        cells.append(
            f'<mxCell id="{eidx}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;html=1;endArrow=block;endFill=1;" edge="1" parent="1" source="{src + 2}" target="{dst + 2}">'
            '<mxGeometry relative="1" as="geometry"/></mxCell>'
        )
    xml = f'<mxfile host="app.diagrams.net"><diagram name="{html.escape(title)}"><mxGraphModel><root>{"".join(cells)}</root></mxGraphModel></diagram></mxfile>'
    (GEN_DIR / f"{name}.drawio").write_text(xml, encoding="utf-8")


def diagram_architecture():
    img, d = base_canvas("SuperFoodie 系统总体架构", "前端、服务、Agent 编排、检索记忆和外部工具的企业级分层")
    nodes = [
        ("React 前端\n表单输入 / 推荐卡片 / PDF 预览", 120, 280, 470, 180),
        ("FastAPI 服务\n路由 / 会话状态 / 报告生成", 760, 280, 470, 180),
        ("Agent 编排器\n意图识别 / 状态机 / 并发任务", 1400, 280, 470, 180),
        ("外部工具层\n小红书 / Tavily / 高德\nMCP 统一接入", 1400, 660, 500, 180),
        ("Graph-RAG + Milvus\n安全校验 / 记忆召回 / 相似去重", 760, 660, 470, 180),
        ("大模型服务\nQwen / DeepSeek\n兼容接口", 120, 660, 470, 180),
        ("本地报告输出\nReportLab PDF + 图片代理", 760, 1060, 520, 180),
    ]
    colors = ["#dbeafe", "#dcfce7", "#fef3c7", "#ffedd5", "#e0f2fe", "#fae8ff", "#ede9fe"]
    outlines = ["#2563eb", "#16a34a", "#d97706", "#f97316", "#0284c7", "#9333ea", "#7c3aed"]
    for node, fill, outline in zip(nodes, colors, outlines):
        box(d, node[1:], node[0], fill, outline)
    for s, e in [((590, 370), (760, 370)), ((1230, 370), (1400, 370)), ((1635, 460), (1635, 660)), ((1400, 750), (1230, 750)), ((760, 750), (590, 750)), ((995, 840), (1020, 1060))]:
        arrow(d, s, e)
    img.save(GEN_DIR / "01_architecture_hd.png", quality=95)
    save_drawio("01_architecture_hd", "系统总体架构", nodes, [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (4, 6)])


def diagram_agent_flow():
    img, d = base_canvas("Agent 交互流程", "更新后的家庭做饭链路：小红书探测与菜名规划并行，详情阶段并发补全")
    nodes = [
        ("用户输入\n口味 / 人数 / 健康限制", 90, 305, 390, 160),
        ("并行任务 A\nXHS Cookie 可用性探测\n短超时，成功即启用", 650, 240, 520, 190),
        ("并行任务 B\nLLM 素菜单规划\n只输出 3 个菜名", 650, 520, 520, 190),
        ("候选列表\nname / menu_slot / keywords\nreason / avoid_reason", 1390, 380, 560, 220),
        ("每个候选并发补详情\nXHS 做法 / LLM 兜底 / Tavily 图片", 840, 850, 760, 200),
        ("合并输出\nXHS 优先，LLM 兜底\n图片由 XHS/Tavily/本地补齐", 840, 1180, 760, 190),
    ]
    fills = ["#dbeafe", "#dcfce7", "#fef3c7", "#ede9fe", "#ffedd5", "#dcfce7"]
    for node, fill in zip(nodes, fills):
        box(d, node[1:], node[0], fill)
    for s, e in [((480, 385), (650, 335)), ((480, 385), (650, 615)), ((1170, 335), (1390, 465)), ((1170, 615), (1390, 465)), ((1670, 600), (1260, 850)), ((1220, 1050), (1220, 1180))]:
        arrow(d, s, e)
    d.text((1650, 960), "单菜失败不拖垮整轮", fill="#475569", font=FT_CAP)
    d.text((1650, 1000), "Tavily 只做图片/来源增强", fill="#475569", font=FT_CAP)
    img.save(GEN_DIR / "02_agent_flow_hd.png", quality=95)
    save_drawio("02_agent_flow_hd", "Agent 交互流程", nodes, [(0, 1), (0, 2), (1, 3), (2, 3), (3, 4), (4, 5)])


def diagram_data_flow():
    img, d = base_canvas("数据流设计", "用户输入、记忆检索、外部内容、结构化推荐与 PDF 输出的数据闭环")
    nodes = [
        ("输入数据\n自然语言、人数、口味、地点、忌口", 110, 280, 520, 180),
        ("上下文构建\n用户画像 + 历史选择 + 本轮约束", 810, 280, 560, 180),
        ("规划数据\n3 个候选菜/餐厅 + 查询关键词", 1550, 280, 560, 180),
        ("内容数据\nXHS 菜谱、Tavily 图片、高德 POI", 1550, 650, 560, 180),
        ("安全与记忆\nGraph-RAG 禁忌规则 + Milvus 去重", 810, 650, 560, 180),
        ("结构化输出\n推荐卡片、步骤、图片、来源、PDF 字段", 810, 1030, 700, 190),
    ]
    for node, fill in zip(nodes, ["#dbeafe", "#dcfce7", "#fef3c7", "#ffedd5", "#e0f2fe", "#ede9fe"]):
        box(d, node[1:], node[0], fill)
    for s, e in [((630, 370), (810, 370)), ((1370, 370), (1550, 370)), ((1830, 460), (1830, 650)), ((1550, 740), (1370, 740)), ((1090, 830), (1160, 1030)), ((1830, 830), (1330, 1030))]:
        arrow(d, s, e)
    d.text((120, 575), "关键原则：LLM 产出必须结构化；外部数据只增强，不让单点失败中断主链路。", fill="#334155", font=FT_CAP)
    img.save(GEN_DIR / "03_data_flow_hd.png", quality=95)
    save_drawio("03_data_flow_hd", "数据流设计", nodes, [(0, 1), (1, 2), (2, 3), (3, 5), (1, 4), (4, 5)])


def diagram_upgrade():
    img, d = base_canvas("系统升级与 AI 能力演进路径", "从可用原型走向可观测、可扩展、可持续演进的 Agent 平台")
    lanes = [
        ("当前版本\n单体 FastAPI + React\n工具可选降级", 120, 300, 500, 190, "#dbeafe"),
        ("下一阶段\n流式卡片展示\n任务队列与缓存\n更细粒度日志", 770, 300, 500, 190, "#dcfce7"),
        ("能力增强\n多模型路由\n评测集 Benchmark\n长期记忆画像", 1420, 300, 500, 190, "#fef3c7"),
        ("工程化平台\n插件化工具市场\n可观测 Agent Trace\n自动化回归评估", 2070, 300, 500, 190, "#ede9fe"),
    ]
    for text, x, y, w, h, fill in lanes:
        box(d, (x, y, w, h), text, fill)
    for s, e in [((620, 395), (770, 395)), ((1270, 395), (1420, 395)), ((1920, 395), (2070, 395))]:
        arrow(d, s, e)
    pillars = [
        ("可扩展架构\n模块边界清晰，工具通过 MCP 接入", 250, 780, 560, 170, "#e0f2fe"),
        ("质量保障\n单元测试、Smoke Test、行为评估并行", 1020, 780, 560, 170, "#ffedd5"),
        ("AI 协作方式\nPrompt、日志、Trace、回放成为新调试入口", 1790, 780, 600, 170, "#fae8ff"),
    ]
    for text, x, y, w, h, fill in pillars:
        box(d, (x, y, w, h), text, fill, fnt=FT_NODE_SMALL)
    img.save(GEN_DIR / "04_upgrade_path_hd.png", quality=95)


def code_card(name: str, title: str, code: str):
    lines = dedent(code).strip("\n").splitlines()
    width = 2200
    height = 190 + len(lines) * 48 + 70
    img = Image.new("RGB", (width, height), "#0f172a")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle([40, 40, width - 40, height - 40], radius=32, fill="#111827", outline="#334155", width=4)
    d.text((90, 72), title, fill="#f8fafc", font=FT_NODE)
    d.text((90, 120), "PlanetB 风格高亮，用于 Word 代码展示", fill="#93c5fd", font=FT_CAP)
    y = 190
    for i, line in enumerate(lines, start=1):
        d.text((90, y), f"{i:>2}", fill="#64748b", font=FT_CODE_SMALL)
        stripped = line.strip()
        if stripped.startswith("#"):
            color = "#94a3b8"
        elif stripped.startswith(("async", "def", "class", "return", "if", "else", "for", "try", "except", "with")):
            color = "#c4b5fd"
        elif "logger" in line or "settings" in line:
            color = "#38bdf8"
        elif "await" in line or "asyncio" in line:
            color = "#fbbf24"
        else:
            color = "#e5e7eb"
        d.text((155, y), line, fill=color, font=FT_CODE)
        y += 48
    img.save(GEN_DIR / name, quality=95)


def generate_assets():
    diagram_architecture()
    diagram_agent_flow()
    diagram_data_flow()
    diagram_upgrade()
    code_card(
        "code_agent_loop_hd.png",
        "Agent 核心循环：规划、工具调用、合并与降级",
        """
        async def run_foodie_agent(request: FoodieRequest) -> FoodieResponse:
            context = build_context(request)
            intent = await classify_intent(context)
            safety = await graph_rag_guard(context)

            if intent == "home_cooking":
                return await run_home_recipe_pipeline(context, safety)
            if intent == "dining_out":
                return await run_dining_pipeline(context, safety)
            return ask_followup_question(context)
        """,
    )
    code_card(
        "code_tool_definition_hd.png",
        "工具定义：外部平台能力统一为可超时任务",
        """
        @tool(timeout=3.0, retry=0)
        async def probe_xhs_available(query: str) -> bool:
            logger.info("xhs_probe.start query=%s", query)
            result = await aione.xhs_search(keyword=query, limit=3)
            return bool(result and result.items)

        @tool(timeout=8.0, retry=1)
        async def search_recipe_images_online(keyword: str) -> ImageBundle:
            return await tavily.search_images(keyword)
        """,
    )
    code_card(
        "code_config_hd.png",
        "配置文件：模型、工具和降级策略可切换",
        """
        LLM_PROVIDER=dashscope
        LLM_MODEL=qwen3.5-plus
        ENABLE_THINKING=false

        XHS_PROBE_TIMEOUT_SECONDS=3
        RECIPE_DETAIL_TIMEOUT_SECONDS=12
        TAVILY_IMAGE_TIMEOUT_SECONDS=6
        FALLBACK_IMAGE_DIR=assets/recipe_placeholders
        """,
    )
    code_card(
        "code_parallel_detail_hd.png",
        "详情阶段：三个候选并发补全，慢任务不阻塞整轮",
        """
        detail_tasks = [
            complete_one_recipe(candidate, xhs_available)
            for candidate in menu_candidates
        ]

        recipes = await asyncio.gather(
            *detail_tasks,
            return_exceptions=False,
        )
        response = build_recommendation_response(recipes)
        """,
    )


def set_shading(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_run(run, size=12, bold=False, color=None, font_name="宋体"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def para(doc: Document, text: str, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    set_run(r, 12)
    return p


def heading(doc: Document, text: str, level: int):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, 16 if level == 1 else 13, True, "111827", "黑体")
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, 12)


def table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        set_shading(c, "DBEAFE")
        for p in c.paragraphs:
            for r in p.runs:
                set_run(r, 10.5, True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    set_run(r, 10.5)
    doc.add_paragraph()
    return t


def image(doc: Document, path: Path, caption: str, width=16.2):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    set_run(r, 9, False, "64748B", "宋体")


def build_report():
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SuperFoodie 超级吃货智能助手")
    set_run(r, 24, True, "111827", "黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CS599 企业级应用软件设计与开发期末大作业报告（增强版）")
    set_run(r, 15, False, "475569", "宋体")
    doc.add_paragraph()
    table(
        doc,
        ["项目", "内容"],
        [
            ["课程方向", "方向一：Agentic AI 原生开发"],
            ["系统定位", "面向家庭做饭与外出就餐的智能饮食推荐 Agent"],
            ["核心能力", "意图识别、Graph-RAG 安全校验、Milvus 记忆召回、MCP 工具调用、PDF 报告生成"],
            ["本版优化", "正文扩写、宋体小四、1.25 倍行距、高清中文图、关键实现代码展示、升级路线补充"],
            ["提交日期", "2026 年 6 月 22 日"],
        ],
    )

    heading(doc, "一、选题背景与设计思想", 1)
    heading(doc, "1.1 问题定义", 2)
    para(doc, "本项目选择“智能饮食推荐”作为应用场景，是因为饮食决策具有明显的多约束特征。用户提出的需求通常并不完整，例如“想吃辣的”“两个人吃点清淡的”“附近找个不贵的店”，这些输入背后同时包含口味、人数、预算、位置、时间、健康限制和信息可信度等因素。系统如果只返回一段泛泛建议，用户仍然需要自己继续搜索、筛选、核对做法和寻找图片，无法形成完整闭环。")
    para(doc, "SuperFoodie 试图解决的问题不是单纯生成菜名，而是把自然语言需求转化为可执行的推荐流程。对于家庭做饭，系统需要给出候选菜、推荐理由、食材、调料、步骤和图片；对于外出就餐，系统需要结合位置、距离、人均、评分和用户偏好给出餐厅建议。最终结果还要能够沉淀为 PDF 报告，方便保存、分享和二次确认。")
    heading(doc, "1.2 现有方案不足", 2)
    para(doc, "传统搜索引擎的不足在于信息分散。搜索“番茄炒蛋做法”会得到大量网页和短视频，但用户需要自行判断是否适合自己的健康限制，也需要手动整理步骤。普通菜谱 App 的不足在于个性化和上下文能力弱，用户每一次都像重新开始，历史偏好无法稳定影响推荐结果。通用聊天机器人虽然能生成较完整的文本，但如果不接入外部工具和结构化校验，很容易出现步骤不可验证、图片缺失、来源不清楚、JSON 截断或对忌口处理不严格的问题。")
    para(doc, "从软件工程角度看，原始的一次性大模型生成方案也存在明显风险。一次让模型生成三个完整菜谱会产生较长输出，容易超过 token 预算或出现 JSON 解析失败；外部平台如果串行调用，总等待时间会被最慢节点拉长；小红书 Cookie 如果失效，整条链路可能长时间等待后才失败。因此，本项目后续把链路拆分为“快速规划菜名”和“并发补全详情”两段。")
    heading(doc, "1.3 项目价值", 2)
    para(doc, "项目价值主要体现在三个层面。第一是用户价值：把“吃什么”这个开放问题变成一组可比较、可确认、可执行的推荐结果。第二是工程价值：把 LLM、向量检索、知识规则、外部 API 和前端展示组织成一个可运行系统，而不是停留在 Prompt Demo。第三是课程价值：项目覆盖 Product Spec、Architecture Spec、API Spec、Agent 状态机、工具定义、测试评估和部署扩展，符合企业级应用软件设计与开发的完整流程。")
    heading(doc, "1.4 技术路线", 2)
    para(doc, "技术路线采用前后端分离和 Agentic Workflow。前端使用 React 承担交互与展示，后端使用 FastAPI 统一管理会话、路由和报告生成，Agent 层负责意图识别、任务编排、外部工具调用与结果合并，数据层结合 Graph-RAG、Milvus 和本地模板形成安全与记忆底座。外部工具通过 MCP 或适配层接入，包括小红书内容搜索、Tavily 图片检索和高德地图餐厅检索。")

    heading(doc, "二、Specs 规格文档", 1)
    para(doc, "Specs 文档是本项目从想法走向工程实现的关键。Product Spec 用来定义用户场景和交互目标，Architecture Spec 用来定义模块拆分和关键数据流，API Spec 用来约束前后端通信格式。三类文档共同构成 SDD 的核心内容，使实现过程有明确边界，也方便后续回归检查。")
    table(
        doc,
        ["规格文档", "核心内容", "在项目中的作用"],
        [
            ["Product Spec", "用户画像、家庭做饭/外出就餐场景、输入字段、结果卡片、PDF 导出", "保证系统围绕用户任务展开，而不是只做模型问答"],
            ["Architecture Spec", "前端、FastAPI、Agent 编排、Graph-RAG、Milvus、MCP、外部 API 的职责划分", "保证系统可维护、可替换、可扩展"],
            ["API Spec", "POST /api/foodie/start、会话状态、推荐结果字段、图片代理和 PDF 导出字段", "保证前后端联调稳定，避免字段随意变化"],
            ["测试与评估 Spec", "功能测试、Agent 行为评估、失败降级、Demo 截图和 Benchmark 口径", "保证最终报告不是只描述功能，而是能证明功能可运行"],
        ],
    )
    para(doc, "在实现过程中，规格文档也暴露了一个重要问题：LLM 输出越长，结构化失败概率越高。因此第二轮实现调整为紧凑菜名规划，第一阶段只返回 name、menu_slot、search_keywords、reason、avoid_reason，食材和步骤放到第二阶段补全。这个变化直接影响了第三章和第四章的架构图与代码展示。")

    heading(doc, "三、系统架构与设计", 1)
    para(doc, "本章重点展示系统架构、Agent 交互流程和数据流设计。文字只做必要说明，核心依靠高清图表达模块关系和数据走向。")
    image(doc, GEN_DIR / "01_architecture_hd.png", "图 1  SuperFoodie 系统总体架构", 16.6)
    para(doc, "总体架构采用清晰分层：React 负责体验，FastAPI 负责服务边界，Agent 编排器负责决策和任务调度，Graph-RAG 与 Milvus 负责安全和记忆，外部工具负责真实世界数据补充，ReportLab 负责最终报告输出。")
    image(doc, GEN_DIR / "02_agent_flow_hd.png", "图 2  Agent 交互流程：XHS 探测与菜名规划并行", 16.6)
    para(doc, "更新后的 Agent 流程把原来的一次长生成拆成两段。第一段并行启动 XHS 可用性探测和 LLM 菜名规划；第二段对三个候选菜并发补全详情。这样既减少了第一轮 JSON 截断风险，也让小红书、LLM 和 Tavily 的耗时可以重叠。")
    image(doc, GEN_DIR / "03_data_flow_hd.png", "图 3  数据流设计：从用户输入到结构化推荐和 PDF", 16.6)
    para(doc, "数据流的核心原则是结构化和可降级。LLM 输出必须符合 schema；小红书和 Tavily 只作为增强数据源，不允许单点失败中断主链路；Graph-RAG 与 Milvus 在推荐前后都参与安全校验与重复控制。")

    heading(doc, "四、关键实现与代码展示", 1)
    para(doc, "关键实现围绕 Agent 核心循环、工具定义、配置文件和 AI IDE 协作展开。代码展示采用 PlanetB 风格高亮图片，避免 Word 中直接粘贴代码造成缩进和字体混乱。")
    image(doc, GEN_DIR / "code_agent_loop_hd.png", "代码展示 1  Agent 核心循环", 16.4)
    para(doc, "Agent 核心循环并不是让模型一次性返回最终答案，而是先构建上下文，再识别意图，随后进入不同业务管道。家庭做饭和外出就餐共享安全校验、会话状态和报告输出，但具体工具链不同。")
    image(doc, GEN_DIR / "code_tool_definition_hd.png", "代码展示 2  工具定义：统一超时和降级策略", 16.4)
    para(doc, "工具定义层把小红书、Tavily、高德等外部能力包装成可超时、可重试、可观测的任务。这样业务代码不直接依赖某个平台的调用细节，也能在 Cookie 失效或网络超时时快速降级。")
    image(doc, GEN_DIR / "code_config_hd.png", "代码展示 3  配置文件：模型与工具策略可切换", 16.4)
    para(doc, "配置文件把模型、思考模式、外部工具超时、本地兜底图片等策略集中管理。本项目在测试 DeepSeek 与 Qwen3.5-plus 时，能够通过模型名和 enable_thinking 参数切换运行方式，避免把实验参数写死在业务代码里。")
    image(doc, GEN_DIR / "code_parallel_detail_hd.png", "代码展示 4  详情阶段并发补全", 16.4)
    para(doc, "AI IDE 在本项目中主要用于快速定位调用链、生成局部重构方案、解释日志和补充测试。传统调试主要看断点和堆栈，而 Agent 应用还需要看 Prompt、工具输入输出、模型返回 JSON、外部接口耗时和降级路径。开发方式因此从“只调代码”扩展为“调代码、调上下文、调工具链和调评测样例”。")

    heading(doc, "五、测试与评估", 1)
    para(doc, "测试与评估分为功能测试、Agent 行为评估、Benchmark 观察和 Demo 截图四部分。由于系统依赖外部平台，测试不能只看一次成功结果，还要验证外部工具不可用时是否能稳定降级。")
    table(
        doc,
        ["评估维度", "测试内容", "结果与说明"],
        [
            ["功能测试", "家庭做饭、外出就餐、PDF 生成、图片代理、推荐卡片展示", "核心流程可运行，页面能展示 3 个候选和详情"],
            ["Agent 行为评估", "意图识别、忌口处理、候选数量、步骤结构化、失败降级", "候选规划改为紧凑 schema 后，JSON 截断风险明显降低"],
            ["Benchmark", "比较一次性完整生成与两段式并发生成的等待结构", "两段式方案更利于流式展示和慢任务隔离，后续可用日志持续量化"],
            ["Demo 截图", "输入页、推荐页、详情页、外出就餐页、PDF 预览", "截图已插入报告，证明系统不是静态设计稿"],
        ],
    )
    para(doc, "功能测试关注系统是否能完成用户任务；Agent 行为评估关注系统是否按预期调用工具、是否遵守安全规则、是否在失败时兜底；Benchmark 关注不同链路设计对等待时间和稳定性的影响。当前版本已经记录 xhs_probe、menu_plan、recipe_detail、tavily_image 等节点日志，后续可以基于这些日志做更细的性能对比。")
    for name, cap in [
        ("bfbcf5f3234eafd9e148a5df6ecf0a35.png", "图 4  首页与输入表单"),
        ("f34f3c16620eb1bdaf65ba465a638323.png", "图 5  用户输入示例"),
        ("5fc144f40f87d8c5100a8c22f27c67ba.png", "图 6  家庭做饭推荐结果"),
        ("4ca37ad51281f08b0afbd7a1740fb780.png", "图 7  菜谱详情与图片展示"),
        ("3d423dd8ae9cbd8e751c5307ffad513e.png", "图 8  外出就餐地图推荐"),
        ("8de757dc6adb74d706ff8151c7650ab1.png", "图 9  PDF 报告预览"),
    ]:
        image(doc, IMG_DIR / name, cap, 15.3)

    heading(doc, "六、系统升级与扩展", 1)
    heading(doc, "6.1 可扩展架构", 2)
    para(doc, "从当前项目状态看，系统已经具备继续扩展的基础：前端、FastAPI、Agent 编排、外部工具和报告输出之间边界相对清晰；小红书、Tavily、高德都可以作为工具插件替换；Graph-RAG 和 Milvus 也可以独立演进。后续如果增加新的平台，例如大众点评、盒马菜谱、营养数据库或学校食堂数据，不需要重写主流程，只需要补充工具适配器和字段归一化逻辑。")
    para(doc, "另一个扩展方向是把当前后端一次性返回结果，升级为卡片级流式返回。系统可以先返回 3 个菜名和推荐理由，再逐个补充步骤、图片和来源。这样用户不会在空白页面等待全部任务结束，也更符合 Agent 多工具并发执行的真实过程。")
    heading(doc, "6.2 下一阶段计划", 2)
    bullet(doc, "第一阶段：完善日志和 Trace，把每个工具调用的输入、输出、耗时、失败原因记录下来，形成可回放的 Agent 调试链路。")
    bullet(doc, "第二阶段：引入任务队列和缓存，对 Tavily 图片、小红书搜索、高德 POI 等结果做短期缓存，减少重复请求和等待时间。")
    bullet(doc, "第三阶段：建设小型 Benchmark 数据集，覆盖辣/清淡/减脂/忌口/多人聚餐/附近餐厅等典型需求，用固定样例持续评估模型和工具链变化。")
    bullet(doc, "第四阶段：把用户确认、收藏、删除和修改结果写入长期记忆，使 Milvus 记忆不只是相似去重，还能真正影响下一次推荐。")
    heading(doc, "6.3 AI 能力演进路径", 2)
    para(doc, "AI 能力演进不应只理解为更换更强模型。对这个项目来说，能力演进包括多模型路由、工具使用策略、上下文压缩、结构化校验和自动评测。简单问题可以交给快模型，复杂问题再调用深度推理模型；图片和来源由检索工具负责，步骤由小红书或 LLM 兜底；所有结果通过 schema 校验后再进入前端。")
    para(doc, "随着 AI 逐渐成为开发主力，系统也应从“人写代码、AI 辅助补全”演进为“人定义目标和约束、AI 生成方案、工程师验证和收敛”。这要求项目本身具备更好的可观测性和可测试性，否则 AI 生成的改动很难判断是否真正改善了系统。")
    image(doc, GEN_DIR / "04_upgrade_path_hd.png", "图 10  系统升级与 AI 能力演进路径", 16.6)

    heading(doc, "七、课程总结", 1)
    heading(doc, "7.1 个人收获", 2)
    para(doc, "通过本项目，我对 Agentic AI 原生开发的理解从“会调用模型”转向“会组织一个由模型、工具、数据和界面共同构成的软件系统”。模型能力很重要，但工程边界更重要。一个可交付系统必须处理异常、延迟、字段兼容、用户体验、测试评估和文档沉淀。")
    para(doc, "另一个收获是认识到规格文档并不是开发前的形式任务，而是后续重构和调试的依据。当我们发现一次性生成完整菜谱容易失败时，能够快速回到 Product Spec 和 Architecture Spec，重新定义第一轮只生成候选菜名，第二轮并发补详情。这种调整不是简单改 Prompt，而是对系统数据流的重新设计。")
    heading(doc, "7.2 工程思维转变", 2)
    para(doc, "AI 逐渐变成开发主力后，工程师的工作重点会发生变化。过去主要关注代码是否能编译、接口是否能跑通；现在还要关注 Prompt 是否稳定、工具调用是否可观测、模型输出是否可验证、失败样例是否能复现。调试方式也从单纯断点调试，扩展到日志追踪、上下文回放、模型对照实验、结构化输出校验和 Benchmark 回归。")
    para(doc, "这种变化要求工程师具备更强的系统判断力。AI 可以快速生成代码和方案，但它并不天然知道项目的真实约束，也不会自动承担质量责任。工程师需要定义清晰目标、拆分可验证任务、判断方案是否符合架构边界，并在测试和日志中确认结果。换句话说，AI 提升了开发速度，但也放大了工程判断的重要性。")
    heading(doc, "7.3 对课程的建议", 2)
    para(doc, "课程可以进一步增加 Agent 调试、评测和工程化部署的内容。很多同学已经能够完成模型调用和页面展示，但真正困难的是如何发现模型失败、如何设计降级、如何建立评测样例、如何让 AI 参与代码开发但不破坏系统质量。如果课程中能加入更多关于 Agent Trace、工具调用日志、结构化输出评估和多模型对比的实践，会更贴近未来的软件开发方式。")
    para(doc, "总体来看，本项目完成了从需求分析、规格设计、架构实现、工具集成、测试评估到报告交付的完整闭环。它不是一个只展示模型能力的 Demo，而是一个尝试把 AI 能力纳入企业级软件工程流程的实践项目。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "附录：工具与生成资产", 1)
    table(
        doc,
        ["资产", "路径"],
        [
            ["增强版 Word", str(OUT_DOC)],
            ["高清中文图与 draw.io 源文件", str(GEN_DIR)],
            ["PlanetB Syntax Highlighter", str(TOOLS_DIR / "planetb-syntax-highlighter")],
            ["ai-drawio", str(TOOLS_DIR / "ai-drawio")],
        ],
    )
    doc.save(OUT_DOC)


def main():
    generate_assets()
    build_report()
    print(OUT_DOC)


if __name__ == "__main__":
    main()
